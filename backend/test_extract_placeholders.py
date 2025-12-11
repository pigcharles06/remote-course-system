"""
Test script to extract all {{placeholder}} from template.docx
This will show all placeholders found in the template and compare with what the code is processing.
"""
from docx import Document
import re
import os

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "resources", "template.docx")

def extract_placeholders_from_paragraphs(doc):
    """Extract placeholders from all paragraphs."""
    placeholders = []
    for para in doc.paragraphs:
        text = para.text
        matches = re.findall(r'\{\{([^}]+)\}\}', text)
        placeholders.extend(matches)
    return placeholders

def extract_placeholders_from_tables(doc):
    """Extract placeholders from all tables."""
    placeholders = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text
                    matches = re.findall(r'\{\{([^}]+)\}\}', text)
                    placeholders.extend(matches)
    return placeholders

def extract_placeholders_from_nested_tables(doc):
    """Extract placeholders from nested tables (tables inside tables)."""
    placeholders = []
    
    def scan_cell_tables(cell):
        """Recursively scan for tables inside a cell."""
        found = []
        for nested_table in cell.tables:
            for row in nested_table.rows:
                for nested_cell in row.cells:
                    for para in nested_cell.paragraphs:
                        text = para.text
                        matches = re.findall(r'\{\{([^}]+)\}\}', text)
                        found.extend(matches)
                    # Recursively check for more nested tables
                    found.extend(scan_cell_tables(nested_cell))
        return found
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                placeholders.extend(scan_cell_tables(cell))
    
    return placeholders

def extract_all_text_runs(doc):
    """Extract placeholders by examining individual text runs (handles split placeholders)."""
    all_text = ""
    
    # From paragraphs
    for para in doc.paragraphs:
        for run in para.runs:
            all_text += run.text
    
    # From tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        all_text += run.text
    
    matches = re.findall(r'\{\{([^}]+)\}\}', all_text)
    return matches

def main():
    print("=" * 60)
    print("Template Placeholder Extractor")
    print("=" * 60)
    print(f"\nTemplate: {TEMPLATE_PATH}")
    
    if not os.path.exists(TEMPLATE_PATH):
        print(f"ERROR: Template not found at {TEMPLATE_PATH}")
        return
    
    doc = Document(TEMPLATE_PATH)
    
    # Extract from different sources
    print("\n" + "-" * 40)
    print("1. Placeholders from PARAGRAPHS:")
    print("-" * 40)
    para_placeholders = extract_placeholders_from_paragraphs(doc)
    for i, p in enumerate(para_placeholders, 1):
        print(f"  {i}. {{{{{p}}}}}")
    print(f"  Total: {len(para_placeholders)}")
    
    print("\n" + "-" * 40)
    print("2. Placeholders from TABLES:")
    print("-" * 40)
    table_placeholders = extract_placeholders_from_tables(doc)
    for i, p in enumerate(table_placeholders, 1):
        print(f"  {i}. {{{{{p}}}}}")
    print(f"  Total: {len(table_placeholders)}")
    
    print("\n" + "-" * 40)
    print("3. Placeholders from NESTED TABLES:")
    print("-" * 40)
    nested_placeholders = extract_placeholders_from_nested_tables(doc)
    for i, p in enumerate(nested_placeholders, 1):
        print(f"  {i}. {{{{{p}}}}}")
    print(f"  Total: {len(nested_placeholders)}")
    
    # Combine all unique
    all_placeholders = set(para_placeholders + table_placeholders + nested_placeholders)
    
    print("\n" + "=" * 60)
    print("SUMMARY - All Unique Placeholders:")
    print("=" * 60)
    sorted_placeholders = sorted(all_placeholders)
    for i, p in enumerate(sorted_placeholders, 1):
        print(f"  {i}. {{{{{p}}}}}")
    print(f"\n  TOTAL UNIQUE: {len(all_placeholders)}")
    
    # Save to file
    output_file = os.path.join(os.path.dirname(__file__), "extracted_placeholders.txt")
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write("All placeholders found in template.docx:\n")
        f.write("=" * 50 + "\n\n")
        for i, p in enumerate(sorted_placeholders, 1):
            f.write(f"{i}. {{{{{p}}}}}\n")
        f.write(f"\nTotal: {len(all_placeholders)}\n")
    print(f"\nSaved to: {output_file}")

if __name__ == "__main__":
    main()
