"""
Word Document Generator using Gemini LLM

This module generates Word documents from form data by:
1. Uploading sample documents to Gemini File API for format reference
2. Using LLM to generate properly formatted values for each placeholder
3. Filling the Word template with generated values
"""

import os
import re
import json
import io
from typing import Dict, List, Any, Optional
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import google.generativeai as genai
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Configure Gemini
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

# Paths
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
RESOURCES_DIR = os.path.join(BASE_DIR, "resources")
TEMPLATE_PATH = os.path.join(RESOURCES_DIR, "template.docx")
SAMPLE_PDF_PATH = os.path.join(RESOURCES_DIR, "sample.pdf")
SAMPLE_DOCX_PATH = os.path.join(RESOURCES_DIR, "sample.docx")


def extract_placeholders(docx_path: str) -> List[str]:
    """Extract all {{placeholder}} patterns from a Word document."""
    placeholders = set()
    doc = Document(docx_path)
    
    def find_in_text(text: str):
        matches = re.findall(r'\{\{(.*?)\}\}', text, re.DOTALL)
        for match in matches:
            # Normalize: replace newlines with spaces, collapse multiple spaces, strip
            normalized = ' '.join(match.split())
            placeholders.add(normalized)
    
    def process_table(table):
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    find_in_text(para.text)
                for nested_table in cell.tables:
                    process_table(nested_table)
    
    for para in doc.paragraphs:
        find_in_text(para.text)
    
    for table in doc.tables:
        process_table(table)
    
    return sorted(placeholders)


def upload_reference_files() -> tuple:
    """Upload sample PDF to Gemini File API."""
    print("Uploading reference files to Gemini...")
    
    pdf_file = None
    
    if os.path.exists(SAMPLE_PDF_PATH):
        pdf_file = genai.upload_file(SAMPLE_PDF_PATH, display_name="sample_format.pdf")
        print(f"  Uploaded PDF: {pdf_file.name}")
    
    return pdf_file


def generate_placeholder_values(
    form_data: Dict[str, Any],
    placeholders: List[str],
    pdf_file
) -> Dict[str, str]:
    """Use Gemini LLM to generate formatted values for each placeholder in batches with retry."""
    
    model_name = os.getenv("MODEL_NAME", "gemini-2.0-flash")
    
    model = genai.GenerativeModel(
        model_name=model_name,
        generation_config={
            "temperature": 0.1,
            "response_mime_type": "application/json",
        },
        system_instruction="""You are a document formatting assistant for a university course application system.
Your task is to take form data and generate properly formatted values that match the style shown in the sample documents.
Pay attention to:
- Checkbox formatting (使用 ■ 表示勾選, □ 表示未勾選)
- Date and semester formatting
- Table data formatting
- Special characters and symbols used in the sample

Output ONLY valid JSON with placeholder keys and their formatted values.
IMPORTANT: You MUST provide a value for EVERY placeholder in the list. Do not skip any."""
    )
    
    # Configure safety settings
    safety_settings = {
        "HARM_CATEGORY_HARASSMENT": "BLOCK_NONE",
        "HARM_CATEGORY_HATE_SPEECH": "BLOCK_NONE",
        "HARM_CATEGORY_SEXUALLY_EXPLICIT": "BLOCK_NONE",
        "HARM_CATEGORY_DANGEROUS_CONTENT": "BLOCK_NONE",
    }
    
    def process_batch(batch: List[str], batch_num: int, total_batches: int, include_pdf: bool = False) -> Dict[str, str]:
        """Process a single batch of placeholders."""
        prompt = f"""
我有一份遠距教學課程申請表單的資料，需要填入Word模板。

請參考附件中的範例文件（PDF），了解正確的格式和符號使用方式。
特別注意：
1. 勾選框使用 ■ 表示已選，□ 表示未選
2. 日期和學期的格式
3. 表格數據的排列方式

## 表單資料 (Form Data):
```json
{json.dumps(form_data, ensure_ascii=False, indent=2)}
```

## 本批次需要填入的欄位 (Placeholders for this batch):
{json.dumps(batch, ensure_ascii=False, indent=2)}

## 任務:
根據表單資料，為每個 placeholder 生成正確格式的值。
- 如果是選擇題，使用 ■ 和 □ 符號
- 週次相關的資料，根據 course_outline_weeks 陣列填入
- 成績評量相關的資料，根據 grading_criteria 陣列填入
- 如果沒有對應資料，根據上下文填入合理的值，不要留空
- 如果需要換行（例如多個選項要分行顯示），請使用 \\n 符號

**重要: 你必須為列表中的每一個 placeholder 都提供值，不能跳過任何一個！**

輸出 JSON 格式，key 必須完全匹配 placeholder 的文字。
"""
        
        content = []
        if pdf_file and include_pdf:
            content.append(pdf_file)
        content.append(prompt)
        
        try:
            response = model.generate_content(content, safety_settings=safety_settings)
            
            if not response.candidates or not response.candidates[0].content.parts:
                print(f"    Batch {batch_num} blocked, skipping...")
                return {}
            
            result_text = response.text.strip()
            
            # Clean up markdown code blocks if present
            if result_text.startswith("```json"):
                result_text = result_text[7:]
            if result_text.startswith("```"):
                result_text = result_text[3:]
            if result_text.endswith("```"):
                result_text = result_text[:-3]
            
            batch_values = json.loads(result_text.strip())
            
            # Validate batch_values is a dictionary
            if not isinstance(batch_values, dict):
                print(f"    Batch {batch_num} returned non-dict type: {type(batch_values)}")
                return {}
            
            return batch_values
            
        except json.JSONDecodeError as e:
            print(f"    Error parsing batch {batch_num}: {e}")
            return {}
        except Exception as e:
            print(f"    Error in batch {batch_num}: {type(e).__name__}: {e}")
            return {}
    
    # Batch processing - split placeholders into groups of 10
    BATCH_SIZE = 10
    MAX_RETRIES = 2
    all_values = {}
    remaining_placeholders = list(placeholders)
    
    for retry_round in range(MAX_RETRIES):
        if not remaining_placeholders:
            break
            
        print(f"\n=== Round {retry_round + 1}/{MAX_RETRIES}: Processing {len(remaining_placeholders)} placeholders ===")
        
        total_batches = (len(remaining_placeholders) + BATCH_SIZE - 1) // BATCH_SIZE
        round_values = {}
        
        for i in range(0, len(remaining_placeholders), BATCH_SIZE):
            batch = remaining_placeholders[i:i + BATCH_SIZE]
            batch_num = i // BATCH_SIZE + 1
            print(f"  Batch {batch_num}/{total_batches}: processing {len(batch)} placeholders...")
            
            # Only include PDF for first batch of first round
            include_pdf = (retry_round == 0 and batch_num == 1)
            batch_values = process_batch(batch, batch_num, total_batches, include_pdf)
            
            print(f"    Got {len(batch_values)} values from batch {batch_num}")
            round_values.update(batch_values)
        
        all_values.update(round_values)
        
        # Check which placeholders are still missing
        remaining_placeholders = [p for p in remaining_placeholders if p not in all_values]
        
        if remaining_placeholders:
            print(f"\n  Still missing {len(remaining_placeholders)} values after round {retry_round + 1}")
        else:
            print(f"\n  All placeholders filled after round {retry_round + 1}!")
    
    if remaining_placeholders:
        print(f"\nWARNING: Still missing {len(remaining_placeholders)} values after {MAX_RETRIES} rounds:")
        for p in remaining_placeholders[:5]:
            print(f"  - {p[:60]}...")
    
    print(f"\nTotal generated values: {len(all_values)} / {len(placeholders)}")
    return all_values


def fill_template(template_path: str, values: Dict[str, str]) -> Document:
    """Fill the Word template with values, handling split placeholders across runs."""
    doc = Document(template_path)
    
    # Debug: Track which placeholders have values
    print(f"[FILL] Received {len(values)} values to fill")
    
    def replace_in_runs(paragraph):
        """Replace placeholders in runs, handling split cases and newlines."""
        if '{{' not in paragraph.text:
            return
        
        # Collect all text and do replacements
        full_text = ''.join(run.text for run in paragraph.runs)
        modified_text = full_text
        
        # Find all placeholders in the text using regex
        placeholder_pattern = r'\{\{(.*?)\}\}'
        
        def replace_match(match):
            raw_key = match.group(1)
            # Normalize the key (same as in extract_placeholders)
            normalized_key = ' '.join(raw_key.split())
            
            # Look up the value
            if normalized_key in values:
                value = values[normalized_key]
                if value is None:
                    return ""
                return str(value)
            else:
                # Key not found, keep original
                return match.group(0)
        
        modified_text = re.sub(placeholder_pattern, replace_match, full_text, flags=re.DOTALL)
        
        # If text was modified, update the runs
        if modified_text != full_text:
            if paragraph.runs:
                first_run = paragraph.runs[0]
                
                # Check if there are line breaks to handle
                if '\n' in modified_text:
                    # Clear first run and handle line breaks
                    first_run.clear()
                    parts = modified_text.split('\n')
                    for i, part in enumerate(parts):
                        first_run.add_text(part)
                        if i < len(parts) - 1:
                            # Add a soft line break (Shift+Enter)
                            br = OxmlElement('w:br')
                            first_run._r.append(br)
                else:
                    # No line breaks, just set text
                    first_run.text = modified_text
                
                # Clear other runs
                for run in paragraph.runs[1:]:
                    run.text = ""
    
    def process_table(table):
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_runs(para)
                for nested_table in cell.tables:
                    process_table(nested_table)
    
    # Process all paragraphs
    for para in doc.paragraphs:
        replace_in_runs(para)
    
    # Process all tables
    for table in doc.tables:
        process_table(table)
    
    return doc


def generate_document(form_data: Dict[str, Any]) -> bytes:
    """
    Main entry point: Generate a filled Word document from form data.
    
    Args:
        form_data: Dictionary containing form field values
        
    Returns:
        bytes: The generated Word document as bytes
    """
    # 1. Extract placeholders from template
    placeholders = extract_placeholders(TEMPLATE_PATH)
    print(f"Found {len(placeholders)} placeholders in template")
    
    # 2. Upload reference files
    pdf_file = upload_reference_files()
    
    # 3. Generate values using LLM
    values = generate_placeholder_values(form_data, placeholders, pdf_file)
    
    # 4. Fill template
    doc = fill_template(TEMPLATE_PATH, values)
    
    # 5. Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()


# Test function
if __name__ == "__main__":
    # Sample form data for testing
    test_form_data = {
        "academic_year": "114",
        "semester": "下學期",
        "main_department": "數位教學中心",
        "co_department": "",
        "degree_level": "碩士班",
        "subject_type": "專業科目",
        "course_name_zh": "測試課程",
        "course_name_en": "Test Course",
        "teacher_name": "測試教師 副教授",
        "permanent_course_id": "TEST001",
        "credits": "3",
        "course_type": "選修",
        "class_count": "1",
        "student_count": "30",
        "language": "否",
        "subtitles": "無字幕",
        "platform": "E3",
        "teaching_method_async_weeks": "8",
        "teaching_method_async_hours": "24",
        "teaching_method_sync_weeks": "4",
        "teaching_method_sync_hours": "12",
        "teaching_method_physical_weeks": "4",
        "teaching_method_physical_hours": "12",
        "teaching_method_total_weeks": "16",
        "teaching_objectives": "本課程旨在...",
        "course_outline_weeks": [
            {"week": 1, "content": "課程介紹", "hours_physical": "3", "hours_async": "", "hours_sync": "", "has_activity": False},
            {"week": 2, "content": "基礎概念", "hours_physical": "", "hours_async": "3", "hours_sync": "", "has_activity": True, "activity_description": "課後作業"},
        ],
        "teaching_activities": ["A.講述", "C.分組報告", "E.議題討論"],
        "grading_criteria": [
            {"category": "平時成績", "percentage": "20", "description": "課堂參與"},
            {"category": "期中報告", "percentage": "30", "description": "小組報告"},
            {"category": "期末報告", "percentage": "50", "description": "專案成果"},
        ],
    }
    
    print("Testing Word document generation...")
    doc_bytes = generate_document(test_form_data)
    
    # Save test output
    output_path = os.path.join(BASE_DIR, "test_output.docx")
    with open(output_path, "wb") as f:
        f.write(doc_bytes)
    print(f"Test document saved to: {output_path}")
