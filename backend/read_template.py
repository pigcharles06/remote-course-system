from docx import Document

def read_docx(file_path):
    try:
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        
        # Also read tables as they often contain form labels
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))
                    
        return "\n".join(full_text)
    except Exception as e:
        return f"Error reading DOCX: {e}"

if __name__ == "__main__":
    file_path = r"c:\Users\user\Desktop\ewant\映竹專案\remote_course_system\backend\template.docx"
    content = read_docx(file_path)
    with open(r"c:\Users\user\Desktop\ewant\映竹專案\remote_course_system\backend\template_content.txt", "w", encoding="utf-8") as f:
        f.write(content)
    print("Saved template content to template_content.txt")
