"""
Analyze a generated Word document to find unreplaced placeholders.
"""
from docx import Document
import re
import sys

def find_unreplaced_placeholders(docx_path: str):
    """Find all {{placeholder}} that are still in the document."""
    doc = Document(docx_path)
    unreplaced = []
    
    def check_paragraph(para, location):
        matches = re.findall(r'\{\{([^}]+)\}\}', para.text)
        for match in matches:
            unreplaced.append({
                'placeholder': f'{{{{{match}}}}}',
                'location': location,
                'full_text': para.text[:100] + '...' if len(para.text) > 100 else para.text
            })
    
    def process_table(table, table_idx):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para_idx, para in enumerate(cell.paragraphs):
                    check_paragraph(para, f'Table {table_idx}, Row {row_idx}, Cell {cell_idx}, Para {para_idx}')
                for nested_idx, nested_table in enumerate(cell.tables):
                    process_table(nested_table, f'{table_idx}.{nested_idx}')
    
    # Check paragraphs
    for para_idx, para in enumerate(doc.paragraphs):
        check_paragraph(para, f'Paragraph {para_idx}')
    
    # Check tables
    for table_idx, table in enumerate(doc.tables):
        process_table(table, table_idx)
    
    return unreplaced

if __name__ == "__main__":
    # Use the file path from command line or default
    file_path = sys.argv[1] if len(sys.argv) > 1 else r"c:\Users\user\Desktop\ewant\映竹專案\新興科技創新與智慧財產權管理_51154e91_教學計畫表.docx"
    
    print(f"Analyzing: {file_path}")
    print("=" * 60)
    
    unreplaced = find_unreplaced_placeholders(file_path)
    
    if unreplaced:
        print(f"\nFound {len(unreplaced)} unreplaced placeholders:\n")
        for i, item in enumerate(unreplaced, 1):
            print(f"{i}. {item['placeholder']}")
            print(f"   Location: {item['location']}")
            print(f"   Context: {item['full_text']}")
            print()
    else:
        print("\n✅ All placeholders have been replaced!")
