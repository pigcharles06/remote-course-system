"""
Diagnostic script to compare template placeholders with generated file.
"""
from docx import Document
import re

# Get template placeholders
template = Document(r'resources/template.docx')
template_phs = set()

def find_placeholders(element, container):
    if hasattr(element, 'paragraphs'):
        for para in element.paragraphs:
            matches = re.findall(r'\{\{([^}]+)\}\}', para.text)
            container.update(m.strip() for m in matches)
    if hasattr(element, 'tables'):
        for table in element.tables:
            for row in table.rows:
                for cell in row.cells:
                    find_placeholders(cell, container)

find_placeholders(template, template_phs)

# Get unreplaced from generated file  
generated = Document(r'c:\Users\user\Desktop\ewant\映竹專案\新興科技創新與智慧財產權管理_c749b5b6_教學計畫表.docx')
unreplaced = set()

find_placeholders(generated, unreplaced)

print(f'Template placeholders: {len(template_phs)}')
print(f'Unreplaced in generated: {len(unreplaced)}')
print()
print('Unreplaced placeholder keys:')
for i, u in enumerate(sorted(unreplaced), 1):
    in_template = "YES" if u in template_phs else "NO"
    print(f'{i}. [Template: {in_template}] {u[:100]}')
